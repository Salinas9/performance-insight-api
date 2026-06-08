from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
DOCX_PATH = OUTPUT_DIR / "Performance_Insight_API_Memoria.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(11, 37, 69)
GRAY = RGBColor(85, 85, 85)
LIGHT_GRAY = RGBColor(242, 244, 247)
BLACK = RGBColor(0, 0, 0)


def set_cell_fill(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)


def apply_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def set_page_geometry(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("Performance Insight API · Migración a Python + FastAPI")
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = GRAY

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Memoria técnica")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = GRAY


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("MEMORIA TÉCNICA")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Performance Insight API en Python")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = BLACK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run(
        "Migración desde ASP.NET, análisis de resultados de JMeter y estructura preparada para Vercel."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = GRAY

    metadata = [
        ("Autora", "Cynthia"),
        ("Fecha", datetime.now().strftime("%d/%m/%Y")),
        ("Tecnologías", "Python 3.12, FastAPI, JMeter 5.6.3, Vercel"),
        ("Estado", "Migración completada y validada en local"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.name = "Calibri"
        label_run.font.size = Pt(11)
        value_run = paragraph.add_run(value)
        value_run.font.name = "Calibri"
        value_run.font.size = Pt(11)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def add_endpoint_table(doc: Document) -> None:
    doc.add_heading("Endpoints principales", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = ["Método", "Ruta", "Descripción"]
    for index, header_text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header_text, bold=True)
        set_cell_fill(cell, "F2F4F7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    rows = [
        ("GET", "/", "Muestra el dashboard HTML generado por el servidor."),
        ("GET", "/api/performance/health", "Comprueba que la API está viva."),
        ("GET", "/api/performance/current-report", "Devuelve el último informe procesado."),
        ("POST", "/api/performance/analyze", "Analiza una carga manual enviada en JSON."),
        ("POST", "/api/performance/dashboard", "Calcula y guarda el dashboard desde muestras JSON."),
        ("POST", "/api/performance/jmeter/upload", "Sube y procesa un .jtl/.csv de JMeter."),
    ]

    for method, route, description in rows:
        row = table.add_row().cells
        set_cell_text(row[0], method, bold=True)
        set_cell_text(row[1], route)
        set_cell_text(row[2], description)
        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_results_table(doc: Document) -> None:
    doc.add_heading("Resultados de las pruebas JMeter", level=2)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = ["Escenario", "Target", "Muestras", "Media", "P95", "Máx", "Errores", "Lectura"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, text, bold=True)
        set_cell_fill(cell, "F2F4F7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    rows = [
        (
            "Smoke local",
            "GET /api/performance/health",
            "6",
            "1.33 ms",
            "2 ms",
            "2 ms",
            "0%",
            "Estado OK; la subida y el parser funcionan.",
        ),
        (
            "Smoke público",
            "dummyjson.com",
            "4",
            "555.5 ms",
            "1640 ms",
            "1640 ms",
            "0%",
            "Se detecta latencia alta en /products?delay=1500.",
        ),
    ]

    for values in rows:
        row = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(row[index], value)
            row[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_document() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    apply_base_styles(doc)
    set_page_geometry(doc)
    add_header_footer(doc)
    add_title_block(doc)

    doc.add_heading("Objetivo del proyecto", level=1)
    add_paragraph(
        doc,
        "El objetivo del proyecto es ofrecer una API que analice resultados de rendimiento, acepte muestras manuales o ficheros de JMeter y muestre un dashboard entendible con métricas técnicas y recomendaciones.",
    )

    doc.add_heading("Migración realizada", level=1)
    add_paragraph(
        doc,
        "La versión original estaba construida en ASP.NET. En esta migración se ha trasladado la API a Python usando FastAPI para facilitar un despliegue posterior en Vercel, manteniendo la idea funcional del proyecto.",
    )
    add_bullet_list(
        doc,
        [
            "Se ha creado una entrada Vercel en api/index.py.",
            "La lógica se ha separado en modelos, rutas y servicios.",
            "Se ha mantenido el parser de JMeter y el dashboard HTML generado en servidor.",
            "Se ha preparado un requirements.txt y un vercel.json para despliegue.",
        ],
    )

    doc.add_heading("Arquitectura de la solución", level=1)
    add_paragraph(
        doc,
        "La aplicación queda organizada de forma modular. El fichero api/index.py expone la app a Vercel. app/main.py crea la instancia FastAPI, app/routes/performance.py publica los endpoints y app/services concentra el análisis, el parser de JMeter, la generación del dashboard y el almacenamiento del último informe.",
    )
    add_bullet_list(
        doc,
        [
            "api/index.py: punto de entrada para Vercel.",
            "app/main.py: creación de la aplicación FastAPI.",
            "app/models/: modelos de entrada, salida y dashboard.",
            "app/services/: análisis, parser CSV/JTL, dashboard HTML y store.",
            "scripts/: utilidades para arrancar FastAPI y lanzar JMeter.",
        ],
    )

    add_endpoint_table(doc)

    doc.add_heading("Flujo de uso", level=1)
    add_paragraph(
        doc,
        "El flujo recomendado consiste en arrancar la API, lanzar una prueba JMeter y subir automáticamente el .jtl para alimentar el dashboard.",
    )
    add_code_block(
        doc,
        [
            "python -m venv .venv",
            ".\\.venv\\Scripts\\python.exe -m pip install -r requirements.txt",
            "powershell -ExecutionPolicy Bypass -File .\\scripts\\run-fastapi.ps1",
            "powershell -ExecutionPolicy Bypass -File .\\scripts\\run-jmeter-smoke.ps1 -Domain dummyjson.com -Threads 1 -Loops 1",
        ],
    )

    add_results_table(doc)

    doc.add_heading("Interpretación de resultados", level=1)
    add_paragraph(
        doc,
        "La prueba local contra /api/performance/health confirma que el nuevo backend Python procesa correctamente las muestras, genera el snapshot y actualiza el dashboard. La prueba contra DummyJSON resulta útil para detectar latencias más realistas y comprobar que el sistema destaca el endpoint más lento y las recomendaciones asociadas.",
    )

    doc.add_heading("Despliegue en Vercel", level=1)
    add_paragraph(
        doc,
        "La estructura del proyecto queda preparada para desplegarse en Vercel gracias a FastAPI. El repositorio incorpora requirements.txt, .python-version, vercel.json y el entrypoint api/index.py. El siguiente paso operativo consiste en subir el repositorio a GitHub y conectarlo a una cuenta de Vercel.",
    )
    add_bullet_list(
        doc,
        [
            "Vercel detecta Python y usa api/index.py como función.",
            "La API puede servirse sin necesidad de mantener ASP.NET.",
            "El estado del último informe está en memoria; para producción convendría persistencia externa.",
        ],
    )

    doc.add_heading("Conclusiones", level=1)
    add_paragraph(
        doc,
        "La migración a Python cumple el objetivo de acercar el proyecto a un despliegue compatible con Vercel. Además, se mantiene el valor académico principal de la práctica: construir una API propia, someterla a pruebas con JMeter y analizar los resultados mediante métricas y visualización.",
    )
    add_paragraph(
        doc,
        "Como mejora futura, sería recomendable desplegarla públicamente, añadir almacenamiento persistente para los informes y guardar histórico de ejecuciones para comparar campañas de rendimiento.",
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
